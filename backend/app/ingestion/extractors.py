"""Format-aware text extraction from file bytes."""

from __future__ import annotations

import io
import re


def _extract_text_utf8(file_bytes: bytes, filename: str) -> str:
    return file_bytes.decode("utf-8")


def _extract_pdf(file_bytes: bytes, filename: str) -> str:
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)

    if not pages:
        raise ValueError(f"No text could be extracted from PDF '{filename}' (may be scanned/image-only)")

    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes, filename: str) -> str:
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))

    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n\n".join(parts)


def _extract_html(file_bytes: bytes, filename: str) -> str:
    from bs4 import BeautifulSoup

    try:
        html = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        html = file_bytes.decode("latin-1")

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup.find_all(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)

    # Collapse multiple blank lines into one
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text


_EXTRACTORS: dict[str, callable] = {
    "text/plain": _extract_text_utf8,
    "text/markdown": _extract_text_utf8,
    "text/csv": _extract_text_utf8,
    "application/json": _extract_text_utf8,
    "application/pdf": _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    "text/html": _extract_html,
}


def extract_text(file_bytes: bytes, mime_type: str, filename: str) -> str:
    """Extract plain text from file bytes based on MIME type.

    Raises ValueError on unsupported type or empty result.
    """
    extractor = _EXTRACTORS.get(mime_type)
    if extractor is None:
        raise ValueError(f"Unsupported file type: {mime_type}")

    text = extractor(file_bytes, filename)

    if not text.strip():
        raise ValueError(f"No text content extracted from '{filename}'")

    return text
